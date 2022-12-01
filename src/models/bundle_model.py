import os
import re
from datetime import date
from os import path
from pathlib import Path
import tempfile
from typing import Any, Iterable

from docx import Document
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTPage, LTTextBoxHorizontal
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from PyPDF2.generic import RectangleObject
from reportlab.pdfgen import canvas

from src.models.exceptions import BundleError


class BundleSectionModel:
    """A class representing the data and underlying subdirectory that forms a section in the bundle."""

    def __init__(self, section_path):
        self.dir_name = os.path.split(section_path)[1]
        self.section, self.name = self._scrape_dir_name()
        self.path = section_path

    def _scrape_dir_name(self):
        dir_name = self.dir_name
        for i, char in enumerate(dir_name):
            if char == '.' and dir_name[i + 1] == ' ':
                return dir_name[:i], dir_name[i+2:]
        raise ValueError(
            f'Section name \'{dir_name}\' is formatted incorrectly.')


class BundleEntryModel:
    """A class representing the data and underlying file that forms an entry in the bundle."""

    def __init__(self, entry_path, tab) -> None:
        self.path = entry_path
        self.tab = tab
        self.file_name = os.path.split(entry_path)[1]

        if not re.fullmatch("[0-9]{4}\.[0-9]{2}\.[0-9]{2} - .+", self.file_name):
            raise ValueError(
                f'Entry name \'{self.file_name}\' is formatted incorrectly.')

        self.date = self._get_date_from_file()
        self.name = self._get_name_from_file()
        self.pag_num = len(PdfReader(self.path).pages)

    def _get_date_from_file(self):
        date_str = self.file_name[0:10]
        year = int(date_str[:4])
        month = int(date_str[5:7])
        day = int(date_str[8:])
        return date(year, month, day)

    def _get_name_from_file(self):
        name = self.file_name[13:]
        name = name[:-4]
        return name


class BundleModel:
    def __init__(self, input_path):
        self._data = self.get_data(input_path)
        self.documents = DocumentsModel(self)
        self.index = self.get_index(input_path)
        self.tmp_dir = None
        self.paths = {
            'input_path': input_path,
        }
        self.name = f'{date.today().year}.{date.today().month}.{date.today().day} {os.path.basename(input_path)}.pdf'

    @property
    def data(self):
        return self._data

    def get_data(self, input_path):
        data = {}
        tab = 1

        for directory in os.listdir(input_path):
            entry_list = []
            sub_dir = os.path.join(input_path, directory)

            if os.path.isdir(sub_dir):
                section = BundleSectionModel(sub_dir)

                for f in os.listdir(sub_dir):
                    entry_list.append(BundleEntryModel(
                        entry_path=os.path.join(sub_dir, f), tab=tab))
                    tab = tab + 1

                data[section] = entry_list

        if data == {}:
            raise BundleError('Bundle contains no sections.')

        return data

    def get_index(self, input_path):
        index_path = os.path.join(input_path + '/index_template.docx')
        if os.path.isfile(index_path):
            return IndexModel(bundle=self, index_path=index_path)
        else:
            return IndexModel(
                bundle=self, index_path='assets/index.docx')

    def get_tmp_dir(self):
        tmp_dir = tempfile.TemporaryDirectory()
        self.paths.update({
            'index_doc_path': os.path.join(tmp_dir.name, "index.docx"),
            'index_pdf_path': os.path.join(tmp_dir.name, "index.pdf"),
            'documents_pdf_path': os.path.join(
                tmp_dir.name, "documents.pdf"),
            'merged_path': os.path.join(tmp_dir.name, "merged.pdf"),
            'pag_input_path': os.path.join(tmp_dir.name, "pag_input.pdf"),
            'pag_output_path': os.path.join(tmp_dir.name, "pag_output.pdf"),
            'link_path': os.path.join(
                self.paths['output_path'], self.name)
        })
        self.tmp_dir = tmp_dir

    def del_tmp_dir(self):
        self.tmp_dir.cleanup()

    def get_entry_paths(self) -> list[str]:
        """Iterate through the bundle data and return a list of paths for each document in the bundle."""

        paths = []
        # for key in self.data:
        #     for entry in self.data[key]:
        #         paths.append(entry.path)
        for key in self._data:
            for entry in self._data[key]:
                paths.append(entry.path)
        return paths

    def get_entries(
        self,
        section: BundleSectionModel = None
    ) -> list[BundleEntryModel]:
        """
        Return a list of BundleEntry objects.

        :param BundleSection section: If empty, function will return every BundleEntry in the Bundle. If defined, only the entries in the specified BundleSection will be returned.
        """

        if section is not None:
            return self.data[section]
        entries = []
        for key in self.data:
            for entry in self.data[key]:
                entries.append(entry)
        return entries


class IndexModel:
    """A class representing the document that will form the index page(s) in the final bundle."""

    def __init__(self, bundle, index_path: str):
        self.bundle = bundle
        self.headings = ('No.', 'Document', 'Date', 'Page No.')
        self.doc = Document(index_path)
        self.table = self.find_table()

    def find_table(self) -> Table:
        """Loop through the Table objects the word document and return the table with the headings as defined in the Index class."""

        def get_cell_text(row_cells: list[_Cell]):
            for cell in row_cells:
                yield cell.text

        for t in self.doc.tables:
            if tuple(get_cell_text(t.row_cells(0))) == self.headings:
                return t

    def add_text(self,
                 cell: _Cell,
                 text: str,
                 alignment: str
                 ):
        """Set the text inside a table cell and format its contents."""

        def format_text():
            para = cell.paragraphs[0]
            para_fmt = para.paragraph_format
            para_fmt.space_before = 76200
            para_fmt.space_after = 76200

            if alignment == 'centre':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell.text = text
        format_text()

    def add_heading(self, section):
        """Add a Header Row in the index table representing a section in the bundle."""
        def set_table_header_bg_color(cell: _Cell):
            """Set background shading for Header Rows."""

            table_cell = cell._tc
            table_cell_properties = table_cell.get_or_add_tcPr()
            cell_shading = OxmlElement('w:shd')
            cell_shading.set(qn('w:fill'), "D5D5D5")
            table_cell_properties.append(cell_shading)
            return cell

        self.table.add_row()
        cells = self.table.rows[-1].cells
        for cell in cells:
            set_table_header_bg_color(cell)

        self.add_text(cells[0], section.section + '.', None)
        self.add_text(cells[1], section.name, None)

    def add_entry(self, entry):
        """Add a Row to the index table representing an entry in the bundle."""

        self.table.add_row()
        cells = self.table.rows[-1].cells
        self.add_text(cells[0], str(entry.tab) + '.', None)
        self.add_text(cells[1], entry.name, None)
        self.add_text(cells[2], str(entry.date), None)

    def input_table_data(self):
        """Add a row in the index table for each section and entry of the bundle."""

        for key in self.bundle.data:
            self.add_heading(key)

            for entry in self.bundle.data[key]:
                self.add_entry(entry)

    def input_pag_nums(self):
        """
        Input the sequential page numbering for each entry in the bundle into the index table.

        :param int offset: This number is the number of pages of the index itself and the page number of the first document in the bundle will follow from this.
        """
        entry_row = 2
        offset = self.get_pag_num()

        for key in self.bundle.data:
            for value in self.bundle.data[key]:
                entry_cell = self.table.rows[entry_row].cells[3]
                if value.pag_num > 1:
                    self.add_text(entry_cell, str(offset + 1) + ' - ' +
                                  str(offset + value.pag_num), 'centre')
                    offset += value.pag_num
                else:
                    offset += value.pag_num
                    self.add_text(entry_cell, str(offset), 'centre')
                entry_row += 1
            entry_row += 1

    def save(self, output_path):
        """Save the index as a word document."""
        self.doc.save(output_path)
        self.doc_path = output_path

    def convert(self, output_path):
        """Convert the index from a word document to a pdf file."""
        convert(self.doc_path, output_path)
        self.pdf_path = output_path

    def get_pag_num(self) -> int:
        """Get the number of pages of the index."""
        return len(PdfReader(self.pdf_path).pages)


class DocumentsModel:
    """A class representing the pdf documents forming the bundle."""

    def __init__(self, bundle):
        self.bundle = bundle
        self.writer = PdfWriter()
        self.reader = None

    def merge_documents(self, output_path: str):
        """Merge the index and documents in the bundle and into a single file."""

        paths = [self.bundle.index.pdf_path]
        paths += self.bundle.get_entry_paths()
        self._pdf_merger(
            paths, output_path)
        self.reader = PdfReader(output_path)

    def _pdf_merger(self, pdf_paths: list[str], output_path: str):
        """Merge the index and documents in the bundle and into a single file."""
        merger = PdfMerger()

        for pdf in pdf_paths:
            merger.append(pdf)

        merger.write(output_path)
        merger.close()

    def paginate(
        self,
        pag_path: str,
        output_path: str
    ):
        """Paginate the bundle."""

        self._pag_gen(pag_path)
        self._apply_pag(pag_path, output_path)

    def _pag_gen(
        self,
        output_path: str
    ):
        """Generate a pdf document with the same page dimensions as the collated bundle."""
        c = canvas.Canvas(output_path)
        pages = len(self.reader.pages)

        for i in range(pages):
            _, y1, x2, y2 = self.reader.getPage(i).mediaBox

            page_num = c.getPageNumber()
            text = str(page_num)
            c.setFont('Helvetica-Bold', 15)
            c.drawString(x2 - 30, y1 + 20, text)
            c.setPageSize((x2, y2))
            c.showPage()
        c.save()

    def _apply_pag(
        self,
        pag_path: str,
        output_path: str
    ):
        """Overlay the collated bundle with the generate page numbers."""

        pag_reader = PdfReader(pag_path)
        page_indices = list(range(0, len(self.reader.pages)))

        for index in page_indices:
            image_page = pag_reader.pages[index]
            content_page = self.reader.pages[index]
            mediabox = content_page.mediabox
            content_page.merge_page(image_page)
            content_page.mediabox = mediabox
            self.writer.add_page(content_page)

        with open(output_path, "wb") as fp:
            self.writer.write(fp)

    def hyperlink(
        self,
        index_path: str,
        bundle_path: str,
        output_path: str
    ):
        """Add hyperlinks from each entry on the index to the respective page in the bundle."""

        bbox_dict = self._get_bbox_dict(index_path)
        line_objects = self._get_line_objects(bbox_dict)
        new_line_objects = self._is_entry(
            line_objects, self.bundle.get_entries())
        self._add_links(bundle_path, output_path,
                        line_objects=new_line_objects)

    def _get_bbox_dict(self, index_path: str):
        """Return a list of Text Objects extracted from the index pdf."""

        index = -1
        text_objects = []
        index_path = Path(index_path).expanduser()

        def show_ltitem_hierarchy(o: Any, index, depth=0):

            if isinstance(o, Iterable):
                for i in o:
                    if isinstance(i, LTPage):
                        index += 1
                    if isinstance(i, LTTextBoxHorizontal):
                        x1, y1, x2, y2 = i.bbox
                        text_object = TextObject(
                            i.get_text().strip(), x1, y1, x2, y2, index)
                        text_objects.append(text_object)
                    show_ltitem_hierarchy(i, index, depth=depth + 1)

        def sort_list(text_objects):
            text_objects.sort(key=lambda x: (x.page, 1/x.y1))
            y1_dict = {}
            for li in text_objects:
                y1 = li.y1

                if y1 in y1_dict:
                    y1_dict[y1].append(li)
                else:
                    y1_dict[y1] = [li]
            return y1_dict

        pages = extract_pages(index_path)

        show_ltitem_hierarchy(pages, index)

        return sort_list(text_objects)

    def _get_line_objects(self, bbox_dict):
        """Return a list of Line Objects from the extracted Text Objects."""

        line_objects = []
        for v in bbox_dict:
            line_object = LineObject(v, bbox_dict[v])
            line_objects.append(line_object)

        return line_objects

    def _is_entry(self, line_objects, bundle_entries):
        """Iterate through the extracted Line Objects and each entry in the bundle.

        Return a list of Line Objects that correspond to a bundle entry."""

        new_bundle_entries = bundle_entries.copy()
        new_line_objects = []

        for line_object in line_objects:
            compiled_text = line_object.compiled_text

            for bundle_entry in bundle_entries:
                if bundle_entry.name in compiled_text:
                    line_object.page = bundle_entry.pag_num
                    new_line_objects.append(line_object)
                    new_bundle_entries.remove(bundle_entry)

        return new_line_objects

    def _add_links(self, bundle_path, output_path, line_objects):
        """For each Line Object, insert a hyperlink to the respective page in the bundle."""

        file = open(bundle_path, 'rb')
        pdf_writer = PdfWriter()
        pdf_reader = PdfReader(file)
        pagdest = self.bundle.index.get_pag_num() + 1

        x1, y1, x2, y2 = pdf_reader.getPage(0).mediaBox
        # print(f'x1, x2: {x1, x2}\ny1, y2: {y1,y2}')

        # add each page in pdf to pdf writer
        num_of_pages = pdf_reader.getNumPages()

        for page in range(num_of_pages):
            current_page = pdf_reader.getPage(page)
            pdf_writer.addPage(current_page)

        for line_object in line_objects:

            text_objects = line_object.text_objects

            for text_object in text_objects:
                pagenum = text_object.page
                x1, y1, x2, y2 = text_object.x1, text_object.y1, text_object.x2, text_object.y2
                # print(x1, y1, x2, y2)

                pdf_writer.addLink(
                    pagenum=pagenum,  # index of the page on which to place the link
                    pagedest=pagdest - 1,  # index of the page to which the link should go
                    # clickable area x1, y1, x2, y2 (starts bottom left corner)
                    rect=RectangleObject([x1, y1, x2, y2]),
                    # border
                    # fit
                )
            pagdest += line_object.page

        with open(path.abspath(output_path), 'wb') as link_pdf:
            pdf_writer.write(link_pdf)
        file.close()


class TextObject:
    """A class containing the extracted text and coordinates for a horizontal text box in a pdf file."""

    def __init__(self, text, x1, y1, x2, y2, page) -> None:
        self.text = text
        self.x1 = x1
        self.y1 = y1
        self.x2 = x2
        self.y2 = y2
        self.page = page


class LineObject:
    """A class forming of each textbox for a given line in a pdf file."""

    def __init__(self, y1, t_obj_list, page=None):
        self.text_objects = t_obj_list
        self.y1 = y1
        compiled_text = ''
        for arg in t_obj_list:
            compiled_text += arg.text
            compiled_text += ' '
        self.compiled_text = compiled_text
        self.page = page
