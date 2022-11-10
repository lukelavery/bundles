from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell, Table
from docx2pdf import convert
from PyPDF2 import PdfReader


class Index:
    """A class representing the document that will form the index page(s) in the final bundle."""

    def __init__(self, path: str):
        self.headings = ('No.', 'Document', 'Date', 'Page No.')
        self.doc = Document(path)
        self.table = self.find_table()

    def find_table(self) -> Table:
        """Loop through the Table objects the word document and return the table with the headings as defined in the Index class."""

        def get_cell_text(row_cells: list[_Cell]):
            for cell in row_cells:
                yield cell.text

        for t in self.doc.tables:
            if tuple(get_cell_text(t.row_cells(0))) == self.headings:
                return t

    def add_text(self,
                 cell: _Cell,
                 text: str,
                 alignment: str
                 ):
        """Set the text inside a table cell and format its contents."""

        def format_text():
            para = cell.paragraphs[0]
            para_fmt = para.paragraph_format
            para_fmt.space_before = 76200
            para_fmt.space_after = 76200

            if alignment == 'centre':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell.text = text
        format_text()

    def add_heading(self, section):
        """Add a Header Row in the index table representing a section in the bundle."""
        def set_table_header_bg_color(cell):
            """Set background shading for Header Rows."""

            tblCell = cell._tc
            tblCellProperties = tblCell.get_or_add_tcPr()
            clShading = OxmlElement('w:shd')
            clShading.set(qn('w:fill'), "D5D5D5")
            tblCellProperties.append(clShading)
            return cell

        self.table.add_row()
        cells = self.table.rows[-1].cells
        for cell in cells:
            set_table_header_bg_color(cell)

        self.add_text(cells[0], section.section + '.', None)
        self.add_text(cells[1], section.name, None)

    def add_entry(self, entry):
        """Add a Row to the index table representing an entry in the bundle."""

        self.table.add_row()
        cells = self.table.rows[-1].cells
        self.add_text(cells[0], str(entry.tab) + '.', None)
        self.add_text(cells[1], entry.name, None)
        self.add_text(cells[2], str(entry.date), None)

    def input_table_data(self, bundle):
        """Add a row in the index table for each section and entry of the bundle."""

        for key in bundle.data:
            self.add_heading(key)

            for entry in bundle.data[key]:
                self.add_entry(entry)

    def input_pag_nums(self,
                       bundle,
                       offset: int
                       ):
        """
        Input the sequential page numbering for each entry in the bundle into the index table.

        :param int offset: This number is the number of pages of the index itself and the page number of the first document in the bundle will follow from this.
        """
        entry_row = 2

        for key in bundle.data:
            for value in bundle.data[key]:
                entry_cell = self.table.rows[entry_row].cells[3]
                if value.pag_num > 1:
                    self.add_text(entry_cell, str(offset + 1) + ' - ' +
                                  str(offset + value.pag_num), 'centre')
                    offset += value.pag_num
                else:
                    offset += value.pag_num
                    self.add_text(entry_cell, str(offset), 'centre')
                entry_row += 1
            entry_row += 1

    def save(self, output_path):
        """Save the index as a word document."""
        self.doc.save(output_path)
        self.doc_path = output_path

    def convert(self, output_path):
        """Convert the index from a word document to a pdf file."""
        convert(self.doc_path, output_path)
        self.pdf_path = output_path

    def get_pag_num(self) -> int:
        """Get the number of pages of the index."""
        return len(PdfReader(self.pdf_path).pages)
