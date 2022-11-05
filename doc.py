from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pdf import get_num_pages
import constants as const


def get_table(doc):
    def get_cell_text(row_cells):
        for c in row_cells:
            yield c.text

    for t in doc.tables:
        if tuple(get_cell_text(t.row_cells(0))) == const.TABLE_HEADINGS:
            return t


def add_text(cell, text, alignment):
    def format_text():
        para = cell.paragraphs[0]
        para_fmt = para.paragraph_format
        para_fmt.space_before = 76200
        para_fmt.space_after = 76200

        if alignment == 'centre':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell.text = text
    format_text()


# def gen_table(self, master):
#     doc = Document(self.paths['index_path'])
#     input_table_data(doc, self.bundle.data)
#     convert(doc, pdf_doc)
#     self.update_pb(master, 33)
#     input_pag_nums(doc, self.bundle.data)


def add_entry(table, entry):
    table.add_row()
    cells = table.rows[-1].cells
    add_text(cells[0], str(entry.tab) + '.', None)
    add_text(cells[1], entry.name, None)
    add_text(cells[2], str(entry.date), None)


def input_table_data(doc, bundle_data):
    table = get_table(doc)

    for key in bundle_data:
        add_heading(table, key)

        for entry in bundle_data[key]:
            add_entry(table, entry)


def gen_table(self, master, bundle, doc, path, index_pdf_path, index_doc_path):
    document = Document(doc)

    word_doc = index_doc_path
    pdf_doc = index_pdf_path

    t = get_table(document)
    index = 1
    num_pages_list = []
    cum_page_list = []

    doc_names = []

    sections = bundle.get_sections()

    for e in sections:
        add_heading(t)
        new_row = t.rows[-1]
        cells = new_row.cells
        add_text(cells[0], e.section + '.', None)
        add_text(cells[1], e.name, None)
        table_data = bundle.get_entries(e)

        for i in range(len(table_data)):
            entry = table_data[i]
            t.add_row()
            new_row = t.rows[-1]
            cells = new_row.cells

            add_text(cells[0], str(entry.tab) + '.', None)
            add_text(cells[1], entry.name, None)
            add_text(cells[2], str(entry.date), None)

            num_pages_list.append(entry.pag_num)
            index = index + 1

            doc_names.append(entry.name)

        document.save(word_doc)

    convert(word_doc, pdf_doc)
    self.pb['value'] = 33
    master.update_idletasks()

    index_pag_num = get_num_pages(pdf_doc)

    index = 0
    entry_row = 2
    cumulative = index_pag_num

    for e in sections:
        for li in bundle.get_entries(e):
            entry_cell = t.rows[entry_row].cells[3]
            old_cumulative = cumulative
            cumulative += num_pages_list[index]
            if num_pages_list[index] > 1:
                add_text(entry_cell, str(old_cumulative + 1) + ' - ' +
                         str(cumulative), 'centre')
                cum_page_list.append(old_cumulative + 1)
            else:
                add_text(entry_cell, str(cumulative), 'centre')
                cum_page_list.append(cumulative)
            entry_row += 1
            index += 1
        entry_row += 1

    document.save(word_doc)
    convert(word_doc, pdf_doc)
    self.pb['value'] = 66
    master.update_idletasks()
    return doc_names, cum_page_list


def input_pag_nums(doc, bundle, offset):
    table = get_table(doc)
    entry_cell = table.rows[entry_row].cells[3]
    entry_row = 2

    for key in bundle.data:
        for value in bundle.data[key]:
            if value.pag_num > 1:
                add_text(entry_cell, str(offset + 1) + ' - ' +
                         str(offset + value.pag_num), 'centre')
                offset += value.pag_num
            else:
                offset += value.pag_num
                add_text(entry_cell, str(offset), 'centre')
            entry_row += 1
        entry_row += 1


# def add_heading(table):
#     def set_table_header_bg_color(cell):
#         """
#         set background shading for Header Rows
#         """
#         tblCell = cell._tc
#         tblCellProperties = tblCell.get_or_add_tcPr()
#         clShading = OxmlElement('w:shd')
#         # Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
#         clShading.set(qn('w:fill'), "D5D5D5")
#         tblCellProperties.append(clShading)
#         return cell

#     table.add_row()
#     cells = table.rows[-1].cells

#     for cell in cells:
#         set_table_header_bg_color(cell)

def add_heading(table, e):
    def set_table_header_bg_color(cell):
        """
        set background shading for Header Rows
        """
        tblCell = cell._tc
        tblCellProperties = tblCell.get_or_add_tcPr()
        clShading = OxmlElement('w:shd')
        clShading.set(qn('w:fill'), "D5D5D5")
        tblCellProperties.append(clShading)
        return cell

    table.add_row()
    cells = table.rows[-1].cells
    for cell in cells:
        set_table_header_bg_color(cell)

    add_text(cells[0], e.section + '.', None)
    add_text(cells[1], e.name, None)
