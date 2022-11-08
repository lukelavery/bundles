from datetime import date
import os
from typing import Any, Iterable
from docx.oxml.ns import qn
from PyPDF2 import PdfReader, PdfFileReader, PdfMerger, PdfWriter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dt import date_to_ymd
from docx.oxml import OxmlElement
from docx2pdf import convert
from reportlab.pdfgen import canvas
from pathlib import Path
from pdfminer.layout import LTTextBoxHorizontal, LTPage
from pdfminer.high_level import extract_pages
from PyPDF2.generic import RectangleObject
from os import path


class TextObject:
    def __init__(self, text, x1, y1, x2, y2, page) -> None:
        self.text = text
        self.x1 = x1
        self.y1 = y1
        self.x2 = x2
        self.y2 = y2
        self.page = page


class LineObject:
    def __init__(self, y1, t_obj_list, page=None):
        self.text_objects = t_obj_list
        self.y1 = y1
        compiled_text = ''
        for arg in t_obj_list:
            compiled_text += arg.text
            compiled_text += ' '
        self.compiled_text = compiled_text
        self.page = page


class BundleEntry:
    def __init__(self, path, tab) -> None:
        self.path = path
        self.tab = tab
        self.file_name = os.path.split(path)[1]
        self.date = self.get_date_from_file()
        self.name = self.get_name_from_file()
        self.pag_num = PdfFileReader(self.path).numPages

    def get_date_from_file(self):
        date_str = self.file_name[0:10]
        year = int(date_str[:4])
        month = int(date_str[5:7])
        day = int(date_str[8:])
        return date(year, month, day)

    def get_name_from_file(self):
        name = self.file_name[13:]
        name = name[:-4]
        return name


class BundleSection:
    def __init__(self, path):
        self.dir_name = os.path.split(path)[1]
        self.section, self.name = self.scrape_dir_name()
        self.path = path

    def scrape_dir_name(self):
        dir_name = self.dir_name
        for i in range(len(dir_name)):
            if dir_name[i] == '.' and dir_name[i + 1] == ' ':
                return dir_name[:i], dir_name[i+2:]


class Bundle:
    def __init__(self, path, index_path):
        self.data = {}
        self.name = date_to_ymd(
            date.today()) + ' ' + os.path.basename(path) + '.pdf'
        tab = 1

        for d in os.listdir(path):
            entry_list = []
            sub_dir = os.path.join(path, d)

            if os.path.isdir(sub_dir):
                section = BundleSection(sub_dir)

                for f in os.listdir(sub_dir):
                    entry_list.append(BundleEntry(
                        path=os.path.join(sub_dir, f), tab=tab))
                    tab = tab + 1

                self.data[section] = entry_list
        self.index = Index(index_path)
        self.documents = Documents(self)

    def get_sections(self):
        return list(self.data.keys())

    def get_entries(self, section=None):
        if section != None:
            return self.data[section]
        else:
            entries = []
            for key in self.data:
                for entry in self.data[key]:
                    entries.append(entry)
            return entries

    def get_paths(self):
        paths = []
        for key in self.data:
            for entry in self.data[key]:
                paths.append(entry.path)
        return paths


class Index:
    def __init__(self, path):
        self.headings = ('No.', 'Document', 'Date', 'Page No.')
        self.doc = Document(path)
        self.table = self.find_table()

    def find_table(self):
        def get_cell_text(row_cells):
            for c in row_cells:
                yield c.text

        for t in self.doc.tables:
            if tuple(get_cell_text(t.row_cells(0))) == self.headings:
                return t

    def add_text(self, cell, text, alignment):
        def format_text():
            para = cell.paragraphs[0]
            para_fmt = para.paragraph_format
            para_fmt.space_before = 76200
            para_fmt.space_after = 76200

            if alignment == 'centre':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell.text = text
        format_text()

    def add_heading(self, e):
        def set_table_header_bg_color(cell):
            """
            set background shading for Header Rows
            """
            tblCell = cell._tc
            tblCellProperties = tblCell.get_or_add_tcPr()
            clShading = OxmlElement('w:shd')
            clShading.set(qn('w:fill'), "D5D5D5")
            tblCellProperties.append(clShading)
            return cell

        self.table.add_row()
        cells = self.table.rows[-1].cells
        for cell in cells:
            set_table_header_bg_color(cell)

        self.add_text(cells[0], e.section + '.', None)
        self.add_text(cells[1], e.name, None)

    def add_entry(self, entry):
        self.table.add_row()
        cells = self.table.rows[-1].cells
        self.add_text(cells[0], str(entry.tab) + '.', None)
        self.add_text(cells[1], entry.name, None)
        self.add_text(cells[2], str(entry.date), None)

    def input_table_data(self, bundle_data):
        for key in bundle_data:
            self.add_heading(key)

            for entry in bundle_data[key]:
                self.add_entry(entry)

    def input_pag_nums(self, bundle, offset):
        entry_row = 2

        for key in bundle.data:
            for value in bundle.data[key]:
                entry_cell = self.table.rows[entry_row].cells[3]
                if value.pag_num > 1:
                    self.add_text(entry_cell, str(offset + 1) + ' - ' +
                                  str(offset + value.pag_num), 'centre')
                    offset += value.pag_num
                else:
                    offset += value.pag_num
                    self.add_text(entry_cell, str(offset), 'centre')
                entry_row += 1
            entry_row += 1

    def save(self, output_path):
        self.doc.save(output_path)
        self.doc_path = output_path

    def convert(self, output_path):
        convert(self.doc_path, output_path)
        self.pdf_path = output_path

    def get_pag_num(self):
        return PdfReader(self.pdf_path).numPages


class Documents:
    def __init__(self, bundle):
        self.bundle = bundle
        self.writer = PdfWriter()

    def merge_documents(self, output_path):
        # fix this
        paths = [self.bundle.index.pdf_path]
        paths += self.bundle.get_paths()
        self.pdf_merger(
            paths, output_path)
        self.reader = PdfReader(output_path)

    def pdf_merger(self, pdf_paths, output_path):
        merger = PdfMerger()

        for pdf in pdf_paths:
            merger.append(pdf)

        merger.write(output_path)
        merger.close()

    def paginate(self, pag_path, output_path):
        def pagGen(self, output_path):
            c = canvas.Canvas(output_path)
            pages = len(self.reader.pages)

            for i in range(pages):
                x1, y1, x2, y2 = self.reader.getPage(i).mediaBox

                page_num = c.getPageNumber()
                text = str(page_num)
                c.setFont('Helvetica-Bold', 15)
                c.drawString(x2 - 30, y1 + 20, text)
                c.setPageSize((x2, y2))
                c.showPage()
            c.save()

        def applyPag(self, pag_path, output_path):

            pag_reader = PdfReader(pag_path)
            page_indices = list(range(0, len(self.reader.pages)))

            for index in page_indices:
                image_page = pag_reader.pages[index]
                content_page = self.reader.pages[index]
                mediabox = content_page.mediabox
                content_page.merge_page(image_page)
                content_page.mediabox = mediabox
                self.writer.add_page(content_page)

            with open(output_path, "wb") as fp:
                self.writer.write(fp)

        pagGen(self, pag_path)
        applyPag(self, pag_path, output_path)

    def hyperlink(self, index_path, bundle_path, output_path):

        def get_bbox_dict(index_path):
            index = -1
            text_objects = []
            path = Path(index_path).expanduser()

            def show_ltitem_hierarchy(o: Any, index, depth=0):

                if isinstance(o, Iterable):
                    for i in o:
                        if isinstance(i, LTPage):
                            index += 1
                        if isinstance(i, LTTextBoxHorizontal):
                            x1, y1, x2, y2 = i.bbox
                            text_object = TextObject(
                                i.get_text().strip(), x1, y1, x2, y2, index)
                            text_objects.append(text_object)
                        show_ltitem_hierarchy(i, index, depth=depth + 1)

            def sort_list(text_objects):
                text_objects.sort(key=lambda x: (x.page, 1/x.y1))
                y1_dict = {}
                for li in text_objects:
                    y1 = li.y1

                    if y1 in y1_dict:
                        y1_dict[y1].append(li)
                    else:
                        y1_dict[y1] = [li]
                return y1_dict

            pages = extract_pages(path)

            show_ltitem_hierarchy(pages, index)

            return sort_list(text_objects)

        def get_line_objects(bbox_dict):
            line_objects = []
            for v in bbox_dict:
                line_object = LineObject(v, bbox_dict[v])
                line_objects.append(line_object)

            return line_objects

        def is_entry(line_objects, bundle_entries):
            new_bundle_entries = bundle_entries.copy()
            new_line_objects = []

            for lo in line_objects:
                compiled_text = lo.compiled_text

                for be in bundle_entries:
                    if be.name in compiled_text:
                        lo.page = be.pag_num
                        new_line_objects.append(lo)
                        new_bundle_entries.remove(be)
                        print(lo.page)

            return new_line_objects

        def add_links(bundle_path, output_path, line_objects):
            file = open(bundle_path, 'rb')
            pdf_writer = PdfWriter()
            pdf_reader = PdfReader(file)
            pagdest = self.bundle.index.get_pag_num() + 1

            x1, y1, x2, y2 = pdf_reader.getPage(0).mediaBox
            # print(f'x1, x2: {x1, x2}\ny1, y2: {y1,y2}')

            # add each page in pdf to pdf writer
            num_of_pages = pdf_reader.getNumPages()

            for page in range(num_of_pages):
                current_page = pdf_reader.getPage(page)
                pdf_writer.addPage(current_page)

            for line_object in line_objects:

                text_objects = line_object.text_objects

                for text_object in text_objects:
                    pagenum = text_object.page
                    x1, y1, x2, y2 = text_object.x1, text_object.y1, text_object.x2, text_object.y2
                    # print(x1, y1, x2, y2)

                    pdf_writer.addLink(
                        pagenum=pagenum,  # index of the page on which to place the link
                        pagedest=pagdest - 1,  # index of the page to which the link should go
                        # clickable area x1, y1, x2, y2 (starts bottom left corner)
                        rect=RectangleObject([x1, y1, x2, y2]),
                        # border
                        # fit
                    )
                print(pagenum, pagdest)
                pagdest += line_object.page

            with open(path.abspath(output_path), 'wb') as link_pdf:
                pdf_writer.write(link_pdf)
            file.close()

        bbox_dict = get_bbox_dict(index_path)
        line_objects = get_line_objects(bbox_dict)
        new_line_objects = is_entry(line_objects, self.bundle.get_entries())
        add_links(bundle_path, output_path, line_objects=new_line_objects)
