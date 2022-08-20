from docx2pdf import convert
import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dir.scraper import scrape_file_names

from index.tbl import add_text
from pdf.pdf import get_num_pages

my_path = 'C:/Users/lukel/Desktop/bundle'


def get_dir_dict(path):
    directory = {}

    for d in os.listdir(path):
        if os.path.isdir(os.path.join(path, d)):
            directory[d] = os.listdir(os.path.join(path, d))

    return directory


def gen_table(dict, doc, path):
    document = Document(doc)
    tables = document.tables

    word_doc = os.path.join(path, 'output.docx')
    pdf_doc = os.path.join(path, 'index.pdf')

    t = tables[3]
    index = 1
    num_pages_list = []
    cum_page_list = []

    doc_names = []

    for e in dict:
        add_heading(t)
        new_row = t.rows[-1]
        cells = new_row.cells
        add_text(cells[0], e[0:2], None)
        add_text(cells[1], e[2:], None)
        table_data = scrape_file_names(dict[e])

        for i in range(len(table_data)):
            (date, name) = table_data[i]
            t.add_row()
            new_row = t.rows[-1]
            cells = new_row.cells

            add_text(cells[0], str(index) + '.', None)
            add_text(cells[1], name, None)
            add_text(cells[2], date, None)

            num_pages_list.append(get_num_pages(
                os.path.join(my_path, e, dict[e][i])))
            index = index + 1

            doc_names.append(name)

        document.save(word_doc)

    convert(word_doc, pdf_doc)

    index_pag_num = get_num_pages(pdf_doc)

    index = 0
    entry_row = 2
    cumulative = index_pag_num

    for e in dict:
        for li in dict[e]:
            entry_cell = t.rows[entry_row].cells[3]
            old_cumulative = cumulative
            cumulative += num_pages_list[index]
            if num_pages_list[index] > 1:
                add_text(entry_cell, str(old_cumulative + 1) + ' - ' +
                         str(cumulative), 'centre')
                cum_page_list.append(old_cumulative + 1)
            else:
                add_text(entry_cell, str(cumulative), 'centre')
                cum_page_list.append(cumulative)
            entry_row += 1
            index += 1
        entry_row += 1

    document.save(word_doc)
    convert(word_doc, pdf_doc)
    return doc_names, cum_page_list


def add_heading(table):
    table.add_row()
    row = table.rows[-1]
    cells = row.cells

    for cell in cells:
        set_table_header_bg_color(cell)


def set_table_header_bg_color(cell):
    """
    set background shading for Header Rows
    """
    tblCell = cell._tc
    tblCellProperties = tblCell.get_or_add_tcPr()
    clShading = OxmlElement('w:shd')
    # Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
    clShading.set(qn('w:fill'), "D5D5D5")
    tblCellProperties.append(clShading)
    return cell


# gen_tabl(dir_dict, "template_index.docx")
